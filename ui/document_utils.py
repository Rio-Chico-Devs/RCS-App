#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
© 2025 RCS - Software Proprietario
Document Utils - Utilità per generazione documenti di produzione
Template aggiornato per matchare il formato richiesto

Version: 1.3.0
Last Updated: 03/10/2025
Author: Sviluppatore antonio

CHANGELOG:
v1.3.0 (03/10/2025):
- Ripristinato codice funzionante
- Gestione flessibile formati materiale (oggetti/dict)
- Layout DOCX semplificato
"""

# type: ignore
# pyright: reportUnknownParameterType=false, reportMissingParameterType=false
# pyright: reportUnknownMemberType=false, reportUnknownVariableType=false
# pyright: reportUnknownArgumentType=false, reportAttributeAccessIssue=false
# pyright: reportUnusedVariable=false
# type: ignore
# pyright: reportUnusedImport=false
# type: ignore  
# pyright: reportPrivateImportUsage=false, reportPrivateUsage=false



import os
from datetime import datetime
from PyQt5.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QMessageBox, QFileDialog

class DocumentUtils:
    """Utilità per generazione documenti di produzione"""
    
    @staticmethod
    def mostra_dialog_formato(parent=None):
        """Mostra dialog per selezione formato documento"""
        dialog = QDialog(parent)
        dialog.setWindowTitle("Genera Documento")
        dialog.setFixedSize(350, 120)

        layout = QVBoxLayout(dialog)
        layout.addWidget(QPushButton("Seleziona formato per il documento:"))

        buttons_layout = QHBoxLayout()
        btn_html = QPushButton("HTML")
        btn_odt = QPushButton("ODT (OpenOffice)")
        btn_cancel = QPushButton("Annulla")

        buttons_layout.addWidget(btn_html)
        buttons_layout.addWidget(btn_odt)
        buttons_layout.addWidget(btn_cancel)
        layout.addLayout(buttons_layout)

        formato_scelto = None

        def on_html():
            nonlocal formato_scelto
            formato_scelto = 'html'
            dialog.accept()

        def on_odt():
            nonlocal formato_scelto
            formato_scelto = 'odt'
            dialog.accept()

        def on_cancel():
            dialog.reject()

        btn_html.clicked.connect(on_html)
        btn_odt.clicked.connect(on_odt)
        btn_cancel.clicked.connect(on_cancel)
        
        if dialog.exec_() == QDialog.Accepted:
            return formato_scelto
        return None
    
    @staticmethod
    def genera_documento_html(preventivo, dati_cliente, parent=None):
        """Genera documento HTML editabile per la produzione"""
        try:
            # Dialog per salvare file
            nome_file = f"SchediaTaglio_{preventivo.codice_preventivo}_{datetime.now().strftime('%Y%m%d')}"
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Salva Scheda di Taglio HTML",
                nome_file,
                "HTML Files (*.html)"
            )
            
            if not file_path:
                return None
                
            # Genera contenuto HTML
            html_content = DocumentUtils._genera_html_template_specifico(preventivo, dati_cliente)
            
            # Salva file
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
                
            # Apri automaticamente
            os.startfile(file_path)
            
            return file_path
            
        except Exception as e:
            if parent:
                QMessageBox.warning(parent, "Errore", f"Errore nella generazione HTML: {e}")
            return None
    
    @staticmethod
    def anteprima_html(preventivo, dati_cliente, parent=None):
        """Mostra anteprima del documento HTML nel browser senza salvarlo"""
        import tempfile
        import webbrowser

        try:
            html_content = DocumentUtils._genera_html_template_specifico(preventivo, dati_cliente)

            # Scrivi in un file temporaneo (delete=False perché il browser deve poterlo aprire)
            with tempfile.NamedTemporaryFile(
                mode='w', suffix='.html', encoding='utf-8', delete=False
            ) as f:
                f.write(html_content)
                temp_path = f.name

            webbrowser.open(f'file://{temp_path}')

        except Exception as e:
            if parent:
                QMessageBox.warning(parent, "Errore", f"Errore nella generazione anteprima:\n{e}")

    @staticmethod
    def genera_documento_docx(preventivo, dati_cliente, parent=None):
        """Genera documento DOCX editabile per la produzione"""
        try:
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                if parent:
                    QMessageBox.information(
                        parent, 
                        "Libreria Mancante",
                        "Per generare DOCX è necessario installare python-docx:\n\npip install python-docx"
                    )
                return None
            
            # Dialog per salvare file
            nome_file = f"SchediaTaglio_{preventivo.codice_preventivo}_{datetime.now().strftime('%Y%m%d')}"
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Salva Scheda di Taglio DOCX",
                nome_file,
                "DOCX Files (*.docx)"
            )
            
            if not file_path:
                return None
                
            # Crea documento
            doc = Document()
            
            # Header centrato
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run("RCS - Scheda di Taglio")
            run.bold = True
            run.font.size = Pt(14)
            
            doc.add_paragraph()  # Spazio
            
            # Informazioni cliente - Tabella 3x3 con ultima riga merged
            info_table = doc.add_table(rows=3, cols=3)
            info_table.style = 'Table Grid'

            # Riga 1: Cliente | Ordine n. | Data
            info_table.cell(0, 0).text = f"Cliente: {dati_cliente.get('nome_cliente', '')}"
            info_table.cell(0, 1).text = f"Ordine n.: {dati_cliente.get('numero_ordine', '')}"
            info_table.cell(0, 2).text = f"Data: {datetime.now().strftime('%d/%m/%Y')}"

            # Riga 2: Codice | Misura | Finitura
            info_table.cell(1, 0).text = f"Codice: {dati_cliente.get('codice', '')}"
            info_table.cell(1, 1).text = f"Misura: {dati_cliente.get('misura', '')}"
            info_table.cell(1, 2).text = f"Finitura: {dati_cliente.get('finitura', '')}"

            # Riga 3: Descrizione (merge delle 3 celle)
            descrizione = dati_cliente.get('oggetto_preventivo', dati_cliente.get('descrizione', ''))
            cell_a = info_table.cell(2, 0)
            cell_c = info_table.cell(2, 2)
            cell_a.merge(cell_c)
            cell_a.text = f"Descrizione: {descrizione}"
            
            doc.add_paragraph()  # Spazio
            
            # Materiali - Layout identico all'HTML
            if hasattr(preventivo, 'materiali') and preventivo.materiali:
                for i, materiale in enumerate(preventivo.materiali):
                    # Gestisci diversi formati di materiale
                    if hasattr(materiale, 'giri'):
                        giri = materiale.giri
                        lunghezza = getattr(materiale, 'lunghezza', 0)
                        sviluppo = getattr(materiale, 'sviluppo', 0)
                        nome = getattr(materiale, 'nome', f'Materiale {i+1}')
                    elif isinstance(materiale, dict):
                        giri = materiale.get('giri', 0)
                        lunghezza = materiale.get('lunghezza', 0)
                        sviluppo = materiale.get('sviluppo', 0)
                        nome = materiale.get('nome', materiale.get('materiale_nome', f'Materiale {i+1}'))
                    else:
                        giri = 1
                        lunghezza = 1000
                        sviluppo = 100
                        nome = f'Materiale {i+1}'
                    
                    # Lunghezza centrata sopra
                    lungh_para = doc.add_paragraph()
                    lungh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    lungh_para.space_after = Pt(2)
                    run = lungh_para.add_run(f"{int(lunghezza)} mm")
                    run.bold = True
                    run.font.size = Pt(10)
                    
                    # Tabella 1x3 per layout orizzontale: Giri | Rettangolo | Sviluppo
                    layout_table = doc.add_table(rows=1, cols=3)
                    # NON applicare stile - tabella senza bordi di default
                    
                    # Celle
                    cell_giri = layout_table.cell(0, 0)
                    cell_rect = layout_table.cell(0, 1)  
                    cell_svil = layout_table.cell(0, 2)
                    
                    # Larghezze
                    layout_table.columns[0].width = Inches(0.5)
                    layout_table.columns[1].width = Inches(5.5)
                    layout_table.columns[2].width = Inches(1.0)
                    
                    # FORZA altezza riga leggermente più alta
                    from docx.oxml.shared import OxmlElement
                    from docx.oxml.ns import qn
                    
                    tr = layout_table.rows[0]._element
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '280')  # Altezza leggermente aumentata (280 twips = ~5mm)
                    trHeight.set(qn('w:hRule'), 'exact')
                    trPr.append(trHeight)
                    
                    # Aggiungi bordi SOLO al rettangolo centrale
                    tcPr = cell_rect._element.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    
                    # Aggiungi tutti e 4 i bordi al rettangolo
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '12')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '000000')
                        tcBorders.append(border)
                    
                    tcPr.append(tcBorders)
                    
                    # Giri (allineata a destra, NESSUN bordo)
                    para_giri = cell_giri.paragraphs[0]
                    para_giri.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para_giri.paragraph_format.space_before = Pt(0)
                    para_giri.paragraph_format.space_after = Pt(0)
                    run_giri = para_giri.add_run(f"G{giri}")
                    run_giri.bold = True
                    run_giri.font.size = Pt(10)
                    
                    # Rettangolo centrale - == a sinistra, nome centrato con spazi
                    para_rect = cell_rect.paragraphs[0]
                    para_rect.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para_rect.paragraph_format.space_before = Pt(0)
                    para_rect.paragraph_format.space_after = Pt(0)
                    
                    # == a sinistra
                    run_eq = para_rect.add_run("==")
                    run_eq.font.size = Pt(8)
                    run_eq.bold = True
                    
                    # Spazi per distanziare il nome (ridotti per mantenere visibilità)
                    para_rect.add_run("               ")  # ~15 spazi
                    
                    # Nome materiale 
                    run_nome = para_rect.add_run(nome)
                    run_nome.bold = True
                    run_nome.font.size = Pt(10)
                    
                    # Sviluppo (allineata a sinistra, NESSUN bordo)
                    para_svil = cell_svil.paragraphs[0]
                    para_svil.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_svil = para_svil.add_run(f"H {int(sviluppo)} mm")
                    run_svil.bold = True
                    run_svil.font.size = Pt(10)
                    
                    doc.add_paragraph()  # Spazio tra materiali
            
            # Tabella operazioni
            doc.add_paragraph()
            ops_table = doc.add_table(rows=6, cols=7)
            ops_table.style = 'Table Grid'
            
            # Header tabella
            headers = ['Operazione', 'Inizio', 'Fine', 'Tempo Tot.', 'Num. Pezzi', 'Data', 'Note']
            for i, header in enumerate(headers):
                cell = ops_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Larghezze colonne (Note più ampia)
            for i, width in enumerate([1.0, 0.6, 0.6, 0.7, 0.7, 0.6, 3.0]):
                for row in ops_table.rows:
                    row.cells[i].width = Inches(width)
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph(f"Documento generato automaticamente dal sistema RCS - {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer.runs[0]
            run.font.size = Pt(8)
            
            # Salva documento
            doc.save(file_path)

            # Apri automaticamente
            os.startfile(file_path)
            
            return file_path
            
        except Exception as e:
            if parent:
                QMessageBox.warning(parent, "Errore", f"Errore nella generazione DOCX: {e}")
            return None
    
    @staticmethod
    def _odt_manifest(svg_paths=None):
        extra = ''
        if svg_paths:
            extra += '  <manifest:file-entry manifest:media-type="" manifest:full-path="Pictures/"/>\n'
            for path in svg_paths:
                extra += f'  <manifest:file-entry manifest:media-type="image/svg+xml" manifest:full-path="{path}"/>\n'
        return (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2">\n'
            '  <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.text" manifest:full-path="/"/>\n'
            '  <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="content.xml"/>\n'
            '  <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="styles.xml"/>\n'
            + extra +
            '</manifest:manifest>'
        )

    @staticmethod
    def _odt_styles():
        return (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<office:document-styles\n'
            '    xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"\n'
            '    xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"\n'
            '    xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0">\n'
            '  <office:styles>\n'
            '    <style:style style:name="Standard" style:family="paragraph" style:class="text"/>\n'
            '  </office:styles>\n'
            '  <office:automatic-styles>\n'
            '    <style:page-layout style:name="pm1">\n'
            '      <style:page-layout-properties fo:page-width="21.001cm" fo:page-height="29.7cm"\n'
            '          style:print-orientation="portrait"\n'
            '          fo:margin-top="1cm" fo:margin-bottom="1cm"\n'
            '          fo:margin-left="1.5cm" fo:margin-right="1.5cm"/>\n'
            '    </style:page-layout>\n'
            '  </office:automatic-styles>\n'
            '  <office:master-styles>\n'
            '    <style:master-page style:name="Standard" style:page-layout-name="pm1"/>\n'
            '  </office:master-styles>\n'
            '</office:document-styles>'
        )

    @staticmethod
    def _odt_content(preventivo, dati_cliente, sc, svg_files=None):
        """Costruisce content.xml ODT senza dipendenze esterne.
        svg_files: lista opzionale a cui vengono aggiunti (path, svg_content) per le coniche.
        """
        def x(v):
            if v is None:
                return ''
            return str(v).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')

        fmap = {'11px': '11pt', '10px': '10pt', '9px': '9pt', '8px': '8pt', '7px': '7pt', '6px': '6pt'}
        ft_info = fmap.get(sc['font_info'], '10pt')
        ft_nome = fmap.get(sc['font_nome'], '11pt')
        ft_giri = fmap.get(sc['font_giri'], '10pt')
        rmap = {'8mm': '0.8cm', '6mm': '0.6cm', '5mm': '0.5cm', '4mm': '0.4cm'}
        rect_h = rmap.get(sc['rect_height'], '0.8cm')
        rect_h_cm = float(rect_h.replace('cm', ''))
        mmap = {'12mm': '1.2cm', '6mm': '0.6cm', '3mm': '0.3cm', '2mm': '0.2cm', '1mm': '0.1cm', '0mm': '0cm'}
        margin_mat = mmap.get(sc['margin_mat'], '1.2cm')
        num_mat = len(preventivo.materiali) if hasattr(preventivo, 'materiali') and preventivo.materiali else 0
        pad  = '0.15cm' if num_mat <= 10 else ('0.08cm' if num_mat <= 17 else '0.04cm')
        cpad = '0.2cm'  if num_mat <= 10 else ('0.12cm' if num_mat <= 17 else '0.08cm')
        npad = '0.1cm'  if num_mat <= 10 else ('0.05cm' if num_mat <= 17 else '0.02cm')

        pad_cm_val = float(pad.replace('cm', ''))

        NS = (
            'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
            'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
            'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
            'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" '
            'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
            'xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0" '
            'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" '
            'xmlns:xlink="http://www.w3.org/1999/xlink"'
        )

        ops_w = ['2.5cm', '1.8cm', '1.8cm', '1.8cm', '1.8cm', '1.8cm', '6cm']
        ops_col_styles = ''.join(
            f'<style:style style:name="TCO{j}" style:family="table-column">'
            f'<style:table-column-properties style:column-width="{w}"/></style:style>'
            for j, w in enumerate(ops_w)
        )

        auto = (
            f'<office:automatic-styles>'
            f'<style:style style:name="PN" style:family="paragraph">'
            f'<style:text-properties fo:font-size="{ft_info}"/></style:style>'
            f'<style:style style:name="PT" style:family="paragraph">'
            f'<style:paragraph-properties fo:text-align="center" fo:margin-bottom="0.3cm"/>'
            f'<style:text-properties fo:font-size="14pt" fo:font-weight="bold"/></style:style>'
            f'<style:style style:name="PC" style:family="paragraph">'
            f'<style:paragraph-properties fo:text-align="center" fo:margin-top="0cm" fo:margin-bottom="0cm"/>'
            f'<style:text-properties fo:font-size="{ft_nome}" fo:font-weight="bold"/></style:style>'
            f'<style:style style:name="PR" style:family="paragraph">'
            f'<style:paragraph-properties fo:text-align="end" fo:margin-top="0cm" fo:margin-bottom="0cm"/>'
            f'<style:text-properties fo:font-size="{ft_giri}" fo:font-weight="bold"/></style:style>'
            f'<style:style style:name="PL" style:family="paragraph">'
            f'<style:paragraph-properties fo:text-align="center" fo:margin-top="{margin_mat}" fo:margin-bottom="0cm"/>'
            f'<style:text-properties fo:font-size="{ft_info}" fo:font-weight="bold"/></style:style>'
            f'<style:style style:name="TI" style:family="table">'
            f'<style:table-properties style:width="17.7cm" table:align="margins"/></style:style>'
            f'<style:style style:name="TCI" style:family="table-column">'
            f'<style:table-column-properties style:column-width="5.9cm"/></style:style>'
            f'<style:style style:name="CB" style:family="table-cell">'
            f'<style:table-cell-properties fo:border="0.05pt solid #000000" fo:padding="{cpad}"/></style:style>'
            f'<style:style style:name="CH" style:family="table-cell">'
            f'<style:table-cell-properties fo:border="0.05pt solid #000000" fo:padding="{cpad}" fo:background-color="#f0f0f0"/></style:style>'
            f'<style:style style:name="TM" style:family="table">'
            f'<style:table-properties style:width="16cm" table:align="center"/></style:style>'
            f'<style:style style:name="TCN" style:family="table-column">'
            f'<style:table-column-properties style:column-width="2cm"/></style:style>'
            f'<style:style style:name="TCW" style:family="table-column">'
            f'<style:table-column-properties style:column-width="12cm"/></style:style>'
            f'<style:style style:name="CMB" style:family="table-cell">'
            f'<style:table-cell-properties fo:border="1.5pt solid #000000" fo:padding="{pad}"/></style:style>'
            f'<style:style style:name="CNB" style:family="table-cell">'
            f'<style:table-cell-properties fo:border="none" fo:padding="{npad}"/></style:style>'
            f'<style:style style:name="CMB_WL" style:family="table-cell">'
            f'<style:table-cell-properties '
            f'fo:border-left="3.5pt solid #000000" fo:border-right="0.75pt solid #000000" '
            f'fo:border-top="1.5pt solid #000000" fo:border-bottom="1.5pt solid #000000" '
            f'fo:padding="{pad}"/></style:style>'
            f'<style:style style:name="CMB_WR" style:family="table-cell">'
            f'<style:table-cell-properties '
            f'fo:border-left="0.75pt solid #000000" fo:border-right="3.5pt solid #000000" '
            f'fo:border-top="1.5pt solid #000000" fo:border-bottom="1.5pt solid #000000" '
            f'fo:padding="{pad}"/></style:style>'
            f'<style:style style:name="RH" style:family="table-row">'
            f'<style:table-row-properties style:row-height="{rect_h}" style:use-optimal-row-height="false"/></style:style>'
            f'<style:style style:name="TCFI" style:family="table-column">'
            f'<style:table-column-properties style:column-width="3.5cm"/></style:style>'
            f'<style:style style:name="TCHL_L" style:family="table-column">'
            f'<style:table-column-properties style:column-width="12.5cm"/></style:style>'
            f'<style:style style:name="RFI" style:family="table-row">'
            f'<style:table-row-properties style:row-height="0.5cm" style:use-optimal-row-height="false"/></style:style>'
            f'<style:style style:name="CFI" style:family="table-cell">'
            f'<style:table-cell-properties fo:border="1pt solid #000000" fo:padding="0.05cm"/></style:style>'
            f'<style:style style:name="THL" style:family="table">'
            f'<style:table-properties style:width="16cm" table:align="center" fo:margin-top="{margin_mat}" fo:margin-bottom="0cm"/></style:style>'
            f'<style:style style:name="TO" style:family="table">'
            f'<style:table-properties style:width="17.7cm" table:align="margins"/></style:style>'
            f'<style:style style:name="DiagLine" style:family="graphic">'
            f'<style:graphic-properties draw:stroke="solid" svg:stroke-color="#000000" svg:stroke-width="0.04cm"'
            f' draw:fill="none" fo:wrap="run-through" style:run-through="foreground"'
            f' draw:wrap-influence-on-position="once-concurrent"/></style:style>'
            f'{ops_col_styles}'
            f'</office:automatic-styles>'
        )

        nom  = x(dati_cliente.get('nome_cliente', ''))
        ord_ = x(dati_cliente.get('numero_ordine', ''))
        dat  = datetime.now().strftime('%d/%m/%Y')
        cod  = x(dati_cliente.get('codice', ''))
        mis  = x(dati_cliente.get('misura', ''))
        fin  = x(dati_cliente.get('finitura', ''))
        desc = x(dati_cliente.get('oggetto_preventivo', dati_cliente.get('descrizione', '')))

        def cell(st, txt, span=None):
            sp = f' table:number-columns-spanned="{span}"' if span else ''
            return f'<table:table-cell table:style-name="{st}"{sp}><text:p text:style-name="PN">{txt}</text:p></table:table-cell>'

        info_table = (
            '<table:table table:name="Info" table:style-name="TI">'
            '<table:table-column table:style-name="TCI"/>'
            '<table:table-column table:style-name="TCI"/>'
            '<table:table-column table:style-name="TCI"/>'
            '<table:table-row>'
            + cell('CB', f'Cliente: {nom}') + cell('CB', f'Ordine n.: {ord_}') + cell('CB', f'Data: {dat}') +
            '</table:table-row>'
            '<table:table-row>'
            + cell('CB', f'Codice: {cod}') + cell('CB', f'Misura: {mis}') + cell('CB', f'Finitura: {fin}') +
            '</table:table-row>'
            '<table:table-row>'
            + cell('CB', f'Descrizione: {desc}', span=3) +
            '<table:covered-table-cell/><table:covered-table-cell/>'
            '</table:table-row>'
            '</table:table>'
        )

        mat_parts = []
        if hasattr(preventivo, 'materiali') and preventivo.materiali:
            for i, materiale in enumerate(preventivo.materiali):
                if hasattr(materiale, 'giri'):
                    giri       = materiale.giri
                    lunghezza  = getattr(materiale, 'lunghezza', 0)
                    sviluppo   = getattr(materiale, 'sviluppo', 0)
                    nome       = x(getattr(materiale, 'nome', f'Materiale {i+1}'))
                    is_conica  = getattr(materiale, 'is_conica', False)
                    con_lato   = getattr(materiale, 'conicita_lato', 'sinistra')
                    con_alt    = getattr(materiale, 'conicita_altezza_mm', 0.0)
                    con_lung   = getattr(materiale, 'conicita_lunghezza_mm', 0.0)
                elif isinstance(materiale, dict):
                    giri       = materiale.get('giri', 0)
                    lunghezza  = materiale.get('lunghezza', 0)
                    sviluppo   = materiale.get('sviluppo', 0)
                    nome       = x(materiale.get('nome', materiale.get('materiale_nome', f'Materiale {i+1}')))
                    is_conica  = materiale.get('is_conica', False)
                    con_lato   = materiale.get('conicita_lato', 'sinistra')
                    con_alt    = materiale.get('conicita_altezza_mm', 0.0)
                    con_lung   = materiale.get('conicita_lunghezza_mm', 0.0)
                else:
                    giri=1; lunghezza=1000; sviluppo=100; nome=f'Materiale {i+1}'
                    is_conica=False; con_lato='sinistra'; con_alt=0.0; con_lung=0.0

                # Calcola taglio_info (misure conica) da mostrare PRIMA della lunghezza
                taglio_info = ''
                if is_conica and con_lung > 0:
                    taglio_info = f'L:{con_lung:.0f}'
                    if con_alt > 0:
                        taglio_info += f' / A:{con_alt:.0f}'
                    taglio_info += ' mm'

                if is_conica and con_lung > 0:
                    # Linea diagonale ODF nativa con y1/y2 corretti per toccare i bordi del rettangolo
                    d_cm = 1.5   # larghezza orizzontale della diagonale
                    cw_cm = 11.7  # larghezza contenuto cella (12cm - padding)
                    h_i = rect_h_cm
                    # y1 negativo per toccare il bordo superiore, y2 per toccare il bordo inferiore
                    diag_y1 = f'-{pad_cm_val:.2f}cm'
                    diag_y2 = f'{h_i - pad_cm_val:.2f}cm'
                    lines_xml = ''
                    if con_lato in ('sinistra', 'entrambi'):
                        lines_xml += (
                            f'<draw:line draw:style-name="DiagLine" text:anchor-type="paragraph"'
                            f' svg:x1="0cm" svg:y1="{diag_y1}" svg:x2="{d_cm}cm" svg:y2="{diag_y2}"'
                            f' draw:z-index="{i * 2}"><text:p/></draw:line>'
                        )
                    if con_lato in ('destra', 'entrambi'):
                        lines_xml += (
                            f'<draw:line draw:style-name="DiagLine" text:anchor-type="paragraph"'
                            f' svg:x1="{cw_cm - d_cm:.1f}cm" svg:y1="{diag_y1}"'
                            f' svg:x2="{cw_cm}cm" svg:y2="{diag_y2}"'
                            f' draw:z-index="{i * 2 + 1}"><text:p/></draw:line>'
                        )
                    center_cell = (
                        f'<table:table-cell table:style-name="CMB">'
                        f'<text:p text:style-name="PC">'
                        + lines_xml +
                        f'==   {nome}'
                        f'</text:p>'
                        f'</table:table-cell>'
                    )
                else:
                    center_cell = (
                        f'<table:table-cell table:style-name="CMB">'
                        f'<text:p text:style-name="PC">==          {nome}</text:p>'
                        f'</table:table-cell>'
                    )

                # header_line: prima cella = misure taglio (se conica), seconda = lunghezza
                header_line = (
                    f'<table:table table:name="HL{i}" table:style-name="THL">'
                    f'<table:table-column table:style-name="TCFI"/>'
                    f'<table:table-column table:style-name="TCHL_L"/>'
                    f'<table:table-row table:style-name="RFI">'
                    f'<table:table-cell table:style-name="CFI">'
                    f'<text:p text:style-name="PN">{taglio_info}</text:p>'
                    f'</table:table-cell>'
                    f'<table:table-cell table:style-name="CNB">'
                    f'<text:p text:style-name="PC">{int(lunghezza)} mm</text:p>'
                    f'</table:table-cell>'
                    f'</table:table-row></table:table>'
                )
                mat_parts.append(
                    header_line +
                    f'<table:table table:name="M{i}" table:style-name="TM">'
                    f'<table:table-column table:style-name="TCN"/>'
                    f'<table:table-column table:style-name="TCW"/>'
                    f'<table:table-column table:style-name="TCN"/>'
                    f'<table:table-row table:style-name="RH">'
                    f'<table:table-cell table:style-name="CNB">'
                    f'<text:p text:style-name="PR">G{giri}</text:p></table:table-cell>'
                    + center_cell +
                    f'<table:table-cell table:style-name="CNB">'
                    f'<text:p text:style-name="PN">H {int(sviluppo)} mm</text:p></table:table-cell>'
                    f'</table:table-row></table:table>'
                )

        hdrs_xml = ''.join(
            f'<table:table-cell table:style-name="CH">'
            f'<text:p text:style-name="PC">{h}</text:p></table:table-cell>'
            for h in ['Operazione', 'Inizio', 'Fine', 'Tempo', 'Pezzi', 'Data', 'Note']
        )
        empty_cells = ''.join(
            '<table:table-cell table:style-name="CB">'
            '<text:p text:style-name="PN"></text:p></table:table-cell>'
            for _ in range(7)
        )
        empty_row   = f'<table:table-row>{empty_cells}</table:table-row>'
        ops_cols    = ''.join(f'<table:table-column table:style-name="TCO{j}"/>' for j in range(7))
        ops_table   = (
            f'<table:table table:name="Ops" table:style-name="TO">{ops_cols}'
            f'<table:table-row>{hdrs_xml}</table:table-row>'
            + empty_row * sc['ops_rows']
            + '</table:table>'
        )

        title_xml = '<text:p text:style-name="PT">RCS - Scheda di Taglio</text:p>' if sc['show_title'] else ''

        return (
            '<?xml version="1.0" encoding="UTF-8"?>'
            f'<office:document-content {NS}>'
            f'{auto}'
            f'<office:body><office:text>'
            f'{title_xml}'
            f'{info_table}'
            + ''.join(mat_parts) +
            '<text:p text:style-name="PN"></text:p>'
            f'{ops_table}'
            '</office:text></office:body>'
            '</office:document-content>'
        )

    @staticmethod
    def genera_documento_odt(preventivo, dati_cliente, parent=None):
        """Genera documento ODT per OpenOffice/LibreOffice (senza dipendenze esterne)"""
        try:
            import zipfile

            # Dialog per salvare file
            nome_file = f"SchediaTaglio_{preventivo.codice_preventivo}_{datetime.now().strftime('%Y%m%d')}"
            file_path, _ = QFileDialog.getSaveFileName(
                parent,
                "Salva Scheda di Taglio ODT",
                nome_file,
                "ODT Files (*.odt)"
            )

            if not file_path:
                return None

            # Genera contenuto ODT con metodi nativi (no dipendenze esterne)
            num_mat = len(preventivo.materiali) if hasattr(preventivo, 'materiali') and preventivo.materiali else 0
            sc = DocumentUtils._calcola_scala(num_mat)
            mi = zipfile.ZipInfo('mimetype')
            mi.compress_type = zipfile.ZIP_STORED
            svg_files = []
            content_xml = DocumentUtils._odt_content(preventivo, dati_cliente, sc, svg_files)
            svg_paths = [path for path, _ in svg_files]
            with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                zf.writestr(mi, 'application/vnd.oasis.opendocument.text')
                zf.writestr('META-INF/manifest.xml', DocumentUtils._odt_manifest(svg_paths))
                zf.writestr('styles.xml', DocumentUtils._odt_styles())
                zf.writestr('content.xml', content_xml)
                for svg_path, svg_content in svg_files:
                    zf.writestr(svg_path, svg_content)

            try:
                import sys
                import subprocess
                if sys.platform == 'win32':
                    os.startfile(file_path)
                else:
                    subprocess.Popen(['xdg-open', file_path])
            except Exception:
                pass

            return file_path

        except Exception as e:
            if parent:
                QMessageBox.warning(parent, "Errore", f"Errore nella generazione ODT: {e}")
            return None

    @staticmethod
    def _calcola_scala(num_materiali):
        """Calcola i parametri di scala in base al numero di materiali.
        Ritorna dict con tutte le misure adattive. Tabella operazioni sempre 5 righe."""
        if num_materiali <= 10:
            return {
                'margin_mat': '1mm', 'rect_height': '8mm', 'font_nome': '11px',
                'font_info': '10px', 'font_giri': '10px', 'top_offset': '-6mm',
                'orient_font': '8px', 'orient_width': '18mm',
                'show_title': True, 'ops_rows': 5, 'ops_height': '8mm',
                'info_margin': '2mm', 'ops_margin_top': '4mm',
            }
        elif num_materiali <= 17:
            return {
                'margin_mat': '2mm', 'rect_height': '6mm', 'font_nome': '10px',
                'font_info': '9px', 'font_giri': '9px', 'top_offset': '-4mm',
                'orient_font': '7px', 'orient_width': '15mm',
                'show_title': False, 'ops_rows': 5, 'ops_height': '6mm',
                'info_margin': '1mm', 'ops_margin_top': '3mm',
            }
        elif num_materiali <= 25:
            return {
                'margin_mat': '1mm', 'rect_height': '5mm', 'font_nome': '9px',
                'font_info': '8px', 'font_giri': '8px', 'top_offset': '-3mm',
                'orient_font': '6px', 'orient_width': '12mm',
                'show_title': False, 'ops_rows': 5, 'ops_height': '5mm',
                'info_margin': '1mm', 'ops_margin_top': '2mm',
            }
        else:  # 26-30
            return {
                'margin_mat': '1mm', 'rect_height': '4mm', 'font_nome': '8px',
                'font_info': '7px', 'font_giri': '7px', 'top_offset': '-2mm',
                'orient_font': '6px', 'orient_width': '10mm',
                'show_title': False, 'ops_rows': 5, 'ops_height': '4mm',
                'info_margin': '1mm', 'ops_margin_top': '2mm',
            }

    @staticmethod
    def _genera_html_template_specifico(preventivo, dati_cliente):
        """Template HTML scalabile - adatta il layout al numero di materiali (1-25)"""

        num_materiali = len(preventivo.materiali) if hasattr(preventivo, 'materiali') and preventivo.materiali else 0
        s = DocumentUtils._calcola_scala(num_materiali)

        # Sezione materiali dettagliata
        materiali_html = ""
        if hasattr(preventivo, 'materiali') and preventivo.materiali:
            for i, materiale in enumerate(preventivo.materiali):
                if hasattr(materiale, 'giri'):
                    giri = materiale.giri
                    lunghezza = getattr(materiale, 'lunghezza', 0)
                    sviluppo = getattr(materiale, 'sviluppo', 0)
                    nome = getattr(materiale, 'nome', f'Materiale {i+1}')
                    is_conica = getattr(materiale, 'is_conica', False)
                    con_lato = getattr(materiale, 'conicita_lato', 'sinistra')
                    con_alt = getattr(materiale, 'conicita_altezza_mm', 0.0)
                    con_lung = getattr(materiale, 'conicita_lunghezza_mm', 0.0)
                elif isinstance(materiale, dict):
                    giri = materiale.get('giri', 0)
                    lunghezza = materiale.get('lunghezza', 0)
                    sviluppo = materiale.get('sviluppo', 0)
                    nome = materiale.get('nome', materiale.get('materiale_nome', f'Materiale {i+1}'))
                    is_conica = materiale.get('is_conica', False)
                    con_lato = materiale.get('conicita_lato', 'sinistra')
                    con_alt = materiale.get('conicita_altezza_mm', 0.0)
                    con_lung = materiale.get('conicita_lunghezza_mm', 0.0)
                else:
                    giri = 1; lunghezza = 1000; sviluppo = 100
                    nome = f'Materiale {i+1}'
                    is_conica = False; con_lato = 'sinistra'; con_alt = 0.0; con_lung = 0.0

                # Diagonale per tela conica (solo una linea nell'angolo, nient'altro)
                diag_svg = ''
                if is_conica and con_lung > 0:
                    d = 16
                    lines = []
                    if con_lato in ('sinistra', 'entrambi'):
                        lines.append(f'<line x1="0" y1="0" x2="{d}" y2="20" stroke="currentColor" stroke-width="1.2"/>')
                    if con_lato in ('destra', 'entrambi'):
                        lines.append(f'<line x1="80" y1="0" x2="{80-d}" y2="20" stroke="currentColor" stroke-width="1.2"/>')
                    diag_svg = (
                        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 80 20" preserveAspectRatio="none" '
                        'style="position:absolute;top:0;left:0;width:100%;height:100%;pointer-events:none;z-index:0;">'
                        + ''.join(lines) + '</svg>'
                    )

                # Calcola taglio_info (misure conica) da mostrare PRIMA della lunghezza
                taglio_info_html = ''
                if is_conica and con_lung > 0:
                    taglio_info_html = f'L:{con_lung:.0f}'
                    if con_alt > 0:
                        taglio_info_html += f' / A:{con_alt:.0f}'
                    taglio_info_html += ' mm'

                figura_html = (
                    f'<div style="width: 80mm; height: {s["rect_height"]}; border: 2px solid #000; display: flex; '
                    f'align-items: center; justify-content: space-between; background: #fff; padding: 0 3mm; '
                    f'position: relative; flex-shrink: 0; overflow: hidden;">'
                    + diag_svg +
                    f'<input type="text" placeholder="Orient." style="width: {s["orient_width"]}; border: none; '
                    f'font-size: {s["orient_font"]}; background: transparent; position: relative; z-index: 1;">'
                    f'<strong style="font-size: {s["font_nome"]}; position: absolute; left: 50%; '
                    f'transform: translateX(-50%); z-index: 1;">{nome}</strong>'
                    f'</div>'
                )

                # Riga 1 sinistra: misure taglio (conica) o campo compilabile (normale)
                if taglio_info_html:
                    left_area_html = f'<strong style="font-size: {s["font_info"]};">{taglio_info_html}</strong>'
                else:
                    left_area_html = f'<input type="text" placeholder="" style="width: 100%; height: 5mm; border: 1.5px solid #000; background: #fff; color: #000; font-size: {s["font_info"]}; padding: 0 1mm; box-sizing: border-box;">'

                # Larghezze colonne layout: [campo] [G] [tela 80mm] [H]
                w_campo = '28mm'
                w_g     = '8mm'
                w_h     = '18mm'

                materiali_html += f"""
                <div style="margin: {s['margin_mat']} auto 0; page-break-inside: avoid; width: fit-content;">
                    <!-- Riga 1: misure taglio (o campo compilabile) a sinistra + lunghezza centrata sulla tela -->
                    <div style="display: flex; align-items: center; margin-bottom: 0;">
                        <div style="width: {w_campo}; flex-shrink: 0;">
                            {left_area_html}
                        </div>
                        <div style="width: {w_g}; flex-shrink: 0;"></div>
                        <div style="width: 80mm; text-align: center; font-size: {s['font_info']}; flex-shrink: 0;">
                            <strong>{lunghezza}mm</strong>
                        </div>
                        <div style="width: {w_h}; flex-shrink: 0;"></div>
                    </div>
                    <!-- Riga 2: spazio campo + G + tela + H -->
                    <div style="display: flex; align-items: center;">
                        <div style="width: {w_campo}; flex-shrink: 0;"></div>
                        <div style="width: {w_g}; text-align: right; font-size: {s['font_giri']}; flex-shrink: 0; padding-right: 1mm;">
                            <strong>G{giri}</strong>
                        </div>
                        {figura_html}
                        <div style="width: {w_h}; font-size: {s['font_giri']}; flex-shrink: 0; padding-left: 2mm;">
                            <strong>H {int(sviluppo)} mm</strong>
                        </div>
                    </div>
                </div>
                """

        # Titolo condizionale
        title_html = ""
        if s['show_title']:
            title_html = '<div class="header"><h2 style="text-align: center; margin: 0; font-size: 14px;">RCS - Scheda di Taglio</h2></div>'

        # Righe tabella operazioni
        ops_rows_html = "".join([
            '<tr>' + ''.join([f'<td style="height: {s["ops_height"]};"><textarea class="editable-area" rows="1"></textarea></td>' for _ in range(7)]) + '</tr>'
            for _ in range(s['ops_rows'])
        ])

        return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Scheda di Taglio - {preventivo.codice_preventivo}</title>
    <style>
        @page {{
            size: A4;
            margin: 10mm;
        }}
        body {{
            font-family: Arial, sans-serif;
            max-width: 180mm;
            margin: 0 auto;
            font-size: {s['font_info']};
            line-height: 1.2;
        }}
        .header {{
            text-align: center;
            margin-bottom: 3mm;
        }}
        .info-section {{
            margin-bottom: {s['info_margin']};
        }}
        .info-label {{
            font-weight: bold;
            display: inline-block;
            width: 25mm;
        }}
        .editable-field {{
            border: none;
            border-bottom: 1px solid #000;
            background: transparent;
            font-size: {s['font_info']};
            padding: 1mm 0;
        }}
        .operations-table {{
            width: 100%;
            max-width: 180mm;
            border-collapse: collapse;
            margin: {s['ops_margin_top']} auto 0;
        }}
        .operations-table th, .operations-table td {{
            border: 1px solid #000;
            padding: 1mm;
            text-align: center;
            height: {s['ops_height']};
        }}
        .operations-table th {{
            background: #f0f0f0;
            font-weight: bold;
            font-size: {s['font_info']};
        }}
        .editable-area {{
            width: 100%;
            border: none;
            background: transparent;
            resize: none;
            font-size: {s['font_info']};
            text-align: center;
        }}
        @media print {{
            body {{ margin: 0; }}
            .editable-field {{
                border-bottom: 1px solid #000 !important;
            }}
        }}
    </style>
</head>
<body>
    {title_html}

    <!-- INFORMAZIONI CLIENTE (3 colonne per riga) -->
    <table style="width: 100%; border-collapse: collapse; margin-bottom: {s['info_margin']};">
        <tr>
            <td style="padding: 2mm; border: 1px solid #000; width: 33%;">
                <span class="info-label">Cliente:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('nome_cliente', '')}" style="width: 28mm;">
            </td>
            <td style="padding: 2mm; border: 1px solid #000; width: 33%;">
                <span class="info-label">Ordine n.:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('numero_ordine', '')}" style="width: 25mm;">
            </td>
            <td style="padding: 2mm; border: 1px solid #000; width: 34%;">
                <span class="info-label">Data:</span>
                <input type="text" class="editable-field" value="{datetime.now().strftime('%d/%m/%Y')}" style="width: 25mm;">
            </td>
        </tr>
        <tr>
            <td style="padding: 2mm; border: 1px solid #000;">
                <span class="info-label">Codice:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('codice', '')}" style="width: 28mm;">
            </td>
            <td style="padding: 2mm; border: 1px solid #000;">
                <span class="info-label">Misura:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('misura', '')}" style="width: 25mm;">
            </td>
            <td style="padding: 2mm; border: 1px solid #000;">
                <span class="info-label">Finitura:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('finitura', '')}" style="width: 25mm;">
            </td>
        </tr>
        <tr>
            <td style="padding: 2mm; border: 1px solid #000;" colspan="3">
                <span class="info-label">Descrizione:</span>
                <input type="text" class="editable-field" value="{dati_cliente.get('oggetto_preventivo', dati_cliente.get('descrizione', ''))}" style="width: calc(100% - 30mm);">
            </td>
        </tr>
    </table>

    <!-- SEZIONE MATERIALI -->
    {materiali_html}

    <!-- TABELLA OPERAZIONI -->
    <div style="page-break-inside: avoid;">
    <table class="operations-table">
        <thead>
            <tr>
                <th style="width: 12%;">Operazione</th>
                <th style="width: 8%;">Inizio</th>
                <th style="width: 8%;">Fine</th>
                <th style="width: 8%;">Tempo</th>
                <th style="width: 8%;">Pezzi</th>
                <th style="width: 8%;">Data</th>
                <th style="width: 48%;">Note</th>
            </tr>
        </thead>
        <tbody>
            {ops_rows_html}
        </tbody>
    </table>
    </div>
</body>
</html>
        """